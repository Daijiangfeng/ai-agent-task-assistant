"""
文档加载器。
按文件扩展名解析 PDF/DOCX/TXT/Markdown，输出统一的 Document 列表。
"""

from __future__ import annotations

from pathlib import Path

from app.config.logging import get_logger
from app.rag.base import Document

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}


class DocumentLoader:
    """
    文档加载器。

    根据文件扩展名分发到对应的解析逻辑：
    - .pdf         -> pypdf 按页解析
    - .docx        -> python-docx 按段落解析
    - .txt/.md     -> 直接读取纯文本
    """

    @classmethod
    def load(cls, file_path: str) -> list[Document]:
        """
        加载并解析文档。

        Args:
            file_path: 文件路径。

        Returns:
            Document 列表，每个 Document 带 source 等元数据。

        Raises:
            FileNotFoundError: 文件不存在。
            ValueError: 不支持的文件类型。
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"不支持的文件类型: {suffix}。"
                f"支持: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        logger.info("加载文档", file_path=file_path, suffix=suffix)

        if suffix == ".pdf":
            return cls._load_pdf(path)
        if suffix == ".docx":
            return cls._load_docx(path)
        return cls._load_text(path)

    @staticmethod
    def _load_text(path: Path) -> list[Document]:
        """读取纯文本文件（txt/md）。"""
        content = path.read_text(encoding="utf-8")
        return [
            Document(
                content=content,
                metadata={"source": str(path), "type": path.suffix.lstrip(".")},
            )
        ]

    @staticmethod
    def _load_pdf(path: Path) -> list[Document]:
        """按页解析 PDF 文件。"""
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        documents: list[Document] = []
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                documents.append(
                    Document(
                        content=text,
                        metadata={
                            "source": str(path),
                            "type": "pdf",
                            "page": page_num + 1,
                        },
                    )
                )
        return documents

    @staticmethod
    def _load_docx(path: Path) -> list[Document]:
        """按段落解析 DOCX 文件，合并为单个 Document。"""
        import docx

        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n".join(paragraphs)
        return [
            Document(
                content=content,
                metadata={"source": str(path), "type": "docx"},
            )
        ]
